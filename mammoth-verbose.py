from sys import argv
from lxml import etree, objectify
from lxml.builder import E
from lxml.builder import ElementMaker
from docx import Document
from docx.shared import Inches, Pt
from docx.text.run import Font, Run
from docx.enum.style import WD_STYLE_TYPE
from docx.dml.color import ColorFormat
import shutil
import argparse
import os.path
import mammoth
import zipfile
import inspect
import copy

# function to check if the input file is valid
def is_valid_file(parser, arg):
  if not os.path.exists(arg):
    parser.error("The file %s does not exist!" % arg)
  else:
    return open(arg, 'r')  # return an open file handle

# defining the program options
parser = argparse.ArgumentParser(description='While using the Mammoth docx converter, add options to preserve source class names and formatting information as attributes.')
parser.add_argument("-i", dest="filename", required=True,
                    help="The name of the file to read.", metavar="FILE",
                    type=lambda x: is_valid_file(parser, x))
parser.add_argument('--map', dest='mapStyles', action='store_true', default=True,
                   help='Create a custom map to preserve the source docx style names as classes in the output HTML. Default is True.')
parser.add_argument('--verbose', dest='preserveFormatting', action='store_true', default=False,
                   help='Preserve any formatting applied to the docx styles as attributes in the output HTML. Default is False.')

args = parser.parse_args()

docxfile = args.filename
fileName = docxfile.name

# an empty dict for our ultimate parsed data
verboseAttrs = {}

# function to read the styles.xml file from within the docx;
# this is used for extracting style names and formatting information.
def getWordStyles(myfile):
  fileName = myfile.name
  filePath = os.path.splitext(myfile.name)[0]
  newName = filePath + ".zip"
  shutil.copyfile(fileName, newName)
  zip = zipfile.ZipFile(newName)
  xml_content = zip.read('word/styles.xml')
  return xml_content

def getWordText(myfile):
  fileName = myfile.name
  filePath = os.path.splitext(myfile.name)[0]
  newName = filePath + ".zip"
  shutil.copyfile(fileName, newName)
  zip = zipfile.ZipFile(newName)
  xml_content = zip.read('word/document.xml')
  return xml_content

# get all the formatting attributes for a style
def getAttrs(element, inputKey="data", inputVal="", inputDict={}):
  attrKey = element.tag
  attrKey = inputKey + "-" + attrKey.split("}").pop()
  attributes = element.attrib
  children = list(element)
  attrVal = ""
  if len(attributes) == 0 and len(children) == 0:
    attrVal = "true"
  elif len(attributes) == 1 and element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"):
    attrVal = element.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val")
  elif len(attributes) > 0:
    for name, value in sorted(element.items()):
      name = name.split("}").pop()
      attrVal = attrVal + "%s:%r;" % (name, value)
  # loop through sub-children and add as attr
  walkChildren(element, attrKey, "", inputDict)

  return attrKey, attrVal

# walk through the element tree for a style
def walkChildren(element, inputKey="data", inputVal="", inputDict={}):
  children = list(element)
  attrKey = ""
  attrVal = ""
  for child in element:
    attrKey, attrVal = getAttrs(child, inputKey, "", inputDict)
    inputDict[attrKey] = attrVal
  return attrKey, attrVal, inputDict

# bringing together getWordStyles, getAttrs, and walkChildren
# to create the final verboseDict of all style names and their
# formatting information.
def getAllStyles(myfile):
  # parse the incoming XML
  source = getWordStyles(myfile)
  root = etree.fromstring(source)

  allStyles = {}
  # get all paragraph styles
  for style in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='paragraph']") :
    styleID = style.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId")
    styleID = styleID.replace("(","").replace(")","")
    # reset the dictionary for this style's children
    allAttr = {}
    # get all child elements of the style
    children = list(style)
    # walk the child tree to collect all elements and their attributes into the dictionary
    attrKey, attrVal, allAttr = walkChildren(style, "data", "", allAttr)
    allAttr['data-w-type'] = 'p'
    # add the style ID to the master dictionary with value = the collected child elements
    allStyles[styleID] = allAttr
        
  # get all character styles
  for style in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type='character']") :
    styleID = style.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}styleId")
    # reset the dictionary for this style's children
    allAttr = {}
    # get all child elements of the style
    children = list(style)
    # walk the child tree to collect all elements and their attributes into the dictionary
    attrKey, attrVal, allAttr = walkChildren(style, "data", "", allAttr)
    allAttr['data-w-type'] = 'r'
    # add the style ID to the master dictionary with value = the collected child elements
    allStyles[styleID] = allAttr

  return allStyles

def getAttr(self):
  for attr in self:
      value = getattr(format, attr, None)
      if value != None:
        print(attr, value)

# run this function before style map and getting style defs
def getDirectFormatting(text, styles):
  source = getWordText(text)
  root = etree.fromstring(source)

  styles_source = getWordText(styles)
  styles_root = etree.fromstring(styles_source)

  newstyles = []
  modcounter = 1
  for para in root.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p") :
    # get all formatting on the P (inside pPr)
    para_format = para.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if para.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle"):
      stylename = para.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle")
      newstyle = stylename + "HEDmod" + str(modcount)
    else:
      newstyle = "HEDmod" + str(modcount)
    # create new style
    # add new style to list
    currstyle = styles_root.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val='" + stylename + "']")
    for child in para_format.iterdescendants():
      inStyle = currstyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + child)
      if inStyle:
        inStyle[]
      else:
        currstyle.append(child)

    currstyle = styles_root.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}style[@{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val='" + stylename + "']")
    for child in currstyle.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr").iterdescendants():
      pass
    if currstyle.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}" + child):
      # if item already exists in parent, set all attributes equal to new values
    else:
      # if item does not exist in parent, append it
    for run in 


<w:p w14:paraId="0EA1567D" w14:textId="4F0362F5" w:rsidR="00637946" w:rsidRDefault="00637946" w:rsidP="00661D3A">
<w:pPr><w:pBdr><w:left w:val="single" w:sz="48" w:space="4" w:color="CCFFCC"/></w:pBdr><w:shd w:val="clear" w:color="auto" w:fill="339966"/><w:jc w:val="center"/><w:pStyle w:val="Text-Standardtx"/></w:pPr>
<w:pPr><w:pBdr><w:left w:val="wave" w:sz="12" w:space="4" w:color="FF6666"/></w:pBdr><w:spacing w:line="480" w:lineRule="auto"/><w:ind w:firstLine="720"/></w:pPr>

<w:style w:type="paragraph" w:customStyle="1" w:styleId="FrontSalesQuotefsq">
  <w:name w:val="Front Sales Quote (fsq)"/>
  <w:basedOn w:val="Normal"/> -- styleID
  <w:rsid w:val="0002472B"/>
  <w:pPr><w:pBdr><w:left w:val="wave" w:sz="12" w:space="4" w:color="FF6666"/></w:pBdr><w:spacing w:line="480" w:lineRule="auto"/><w:ind w:firstLine="720"/></w:pPr>
  <w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial" w:cs="Arial"/></w:rPr>
</w:style>

# def getDirectFormatting(document):
#   #fobj = open(myfile,'rb')
#   #document = Document(fobj)
#   paragraphs = document.paragraphs
#   styles = document.styles
#   modcount = 1
#   for para in paragraphs:
#     parastyles = {}
#     actualparastyles = {}
#     # Get paragraph formatting
#     format = para.paragraph_format
#     sublist = [a for a in dir(format) if not a.startswith('_') and a != 'element' and a != 'tab_stops']
#     parastyles['alignment'] = para.paragraph_format.alignment
#     #Element is assumed to be w:p
#     #parastyles['element'] = para.paragraph_format.element
#     parastyles['first_line_indent'] = para.paragraph_format.first_line_indent
#     parastyles['keep_together'] = para.paragraph_format.keep_together
#     parastyles['keep_with_next'] = para.paragraph_format.keep_with_next
#     parastyles['left_indent'] = para.paragraph_format.left_indent
#     parastyles['line_spacing'] = para.paragraph_format.line_spacing
#     parastyles['line_spacing_rule'] = para.paragraph_format.line_spacing_rule
#     parastyles['page_break_before'] = para.paragraph_format.page_break_before
#     #parastyles['part'] = para.paragraph_format.part
#     parastyles['right_indent'] = para.paragraph_format.right_indent
#     parastyles['space_after'] = para.paragraph_format.space_after
#     parastyles['space_before'] = para.paragraph_format.space_before
#     #parastyles['tab_stops'] = para.paragraph_format.tab_stops
#     parastyles['widow_control'] = para.paragraph_format.widow_control
#     for key,val in parastyles.items():
#       if val != None:
#         actualparastyles[key] = val
#         # TO DO: do something with paragraph formatting
#     if bool(actualparastyles) == True:
#       if para.style.name != None:
#         newstyle = para.style.name + "HEDmod" + str(modcount)
#         style = styles.add_style(newstyle, WD_STYLE_TYPE.PARAGRAPH)
#         style.base_style = styles[para.style.name]
#       else:
#         newstyle = "HEDmod" + str(modcount)
#         style = styles.add_style(newstyle, WD_STYLE_TYPE.PARAGRAPH)
#       paragraph_format = style.paragraph_format
#       for key,val in actualparastyles.items():
#         setattr(paragraph_format, key, val)
#       para.style = newstyle
#       modcount +=1
#       #style = styles.add_style(newstyle, WD_STYLE_TYPE.PARAGRAPH)
#     runs = para.runs
#     for run in runs:
#       font = run.font
#       charstyles = {}
#       actualcharstyles = {}
#       sublist = [a for a in dir(font) if not a.startswith('_') and a != 'element']
#       charstyles['all_caps'] = run.font.all_caps
#       charstyles['bold'] = run.font.bold
#       charstyles['color'] = run.font.color.rgb
#       charstyles['complex_script'] = run.font.complex_script
#       charstyles['cs_bold'] = run.font.cs_bold
#       charstyles['cs_italic'] = run.font.cs_italic
#       charstyles['double_strike'] = run.font.double_strike
#       charstyles['emboss'] = run.font.emboss
#       charstyles['hidden'] = run.font.hidden
#       charstyles['highlight_color'] = run.font.highlight_color
#       charstyles['imprint'] = run.font.imprint
#       charstyles['italic'] = run.font.italic
#       charstyles['math'] = run.font.math
#       charstyles['name'] = run.font.name
#       charstyles['no_proof'] = run.font.no_proof
#       charstyles['outline'] = run.font.outline
#       #charstyles['part'] = run.font.part
#       charstyles['rtl'] = run.font.rtl
#       charstyles['shadow'] = run.font.shadow
#       charstyles['size'] = run.font.size
#       charstyles['small_caps'] = run.font.small_caps
#       charstyles['snap_to_grid'] = run.font.snap_to_grid
#       charstyles['spec_vanish'] = run.font.spec_vanish
#       charstyles['strike'] = run.font.strike
#       charstyles['subscript'] = run.font.subscript
#       charstyles['superscript'] = run.font.superscript
#       charstyles['underline'] = run.font.underline
#       charstyles['web_hidden'] = run.font.web_hidden
#       for key,val in charstyles.items():
#         if val != None:
#           actualcharstyles[key] = val
#       if bool(actualcharstyles) == True:
#         if run.style.name != None and run.style.name != 'Default Paragraph Font':
#           newstyle = run.style.name + "HEDmod" + str(modcount)
#           style = styles.add_style(newstyle, WD_STYLE_TYPE.CHARACTER)
#           style.base_style = styles[run.style.name]
#         else:
#           newstyle = "HEDmod" + str(modcount)
#           style = styles.add_style(newstyle, WD_STYLE_TYPE.CHARACTER)
#         for key,val in actualcharstyles.items():
#           if key == 'color':
#             style.font.color.rgb = run.font.color.rgb
#             style.font.color.theme_color = run.font.color.theme_color
#           else:
#             setattr(style.font, key, val)
#         run.style = newstyle
#         modcount +=1
#   return document

# add the formatting info back to the HTML as attributes on each element
def addAttrs(html, myDict):
  root = etree.HTML(html)
  for style, vals in myDict.items():
    for para in root.findall(".//p[@class='" + style + "']"):
      for key, val in vals.items():
        para.attrib[key] = val
  newHTML = etree.tostring(root, encoding="UTF-8", standalone=True, xml_declaration=True)
  return newHTML

def sanitizeHTML(html):
  root = etree.HTML(html)
  newHTML = etree.tostring(root, encoding="UTF-8", standalone=True, xml_declaration=True)
  return newHTML

fobj = open(fileName,'rb')
document = Document(fobj)

document = getDirectFormatting(document)

document.save('tmp.docx')
newfile = open('tmp.docx','rb')

#verboseAttrs = getAllStyles(docxfile)
verboseAttrs = getAllStyles(newfile)

# create the style map if requested
if args.mapStyles == True:
  style_map = '"""'
  for style, vals in verboseAttrs.items():
    sourceName = vals['data-name']
    destName = style
    # mapping paragraphs
    if vals['data-w-type'] == 'p':
      thisMap = "p[style-name='" + sourceName + "'] => p." + destName + ":fresh"
    # mapping runs
    else:
      thisMap = "r[style-name='" + sourceName + "'] => span." + destName
    # write this map to the map file
    style_map = style_map + "\n" + thisMap

  style_map = style_map + '\n"""'

# convert with mammoth
result = mammoth.convert_to_html(newfile, style_map=style_map)
#result = mammoth.convert_to_html(fobj, style_map=style_map)
html = result.value
messages = result.messages

# add the verbose attributes to the output HTML if requested
if args.preserveFormatting == True:
  html = addAttrs(html, verboseAttrs)
else:
  html = sanitizeHTML(html)

# write to a new HTML document
output = open('output.html', 'w')
output.write(str(html))
output.close()